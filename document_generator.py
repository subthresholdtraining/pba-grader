"""
Document Generator Module
Generates Word documents with grading feedback in Amanda's format.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from io import BytesIO
from datetime import datetime
from typing import Dict, List, Tuple

from grading_logic import GradeResult


# Amanda's color scheme
PURPLE = RGBColor(0x77, 0x56, 0xA7)  # #7756A7 - Purple for question headings
LABEL_RED = RGBColor(0xE7, 0x4E, 0x4E)  # #E74E4E - Red for labels (Student Name, etc.)
RED = RGBColor(255, 0, 0)  # Red for INCORRECT
GREEN = RGBColor(0, 128, 0)  # Green for CORRECT

# Font settings
FONT_NAME = 'Lato'
TITLE_SIZE = Pt(18)
BODY_SIZE = Pt(11)


def create_grading_document(
    student_name: str,
    submission_date: str,
    reviewer_name: str,
    answers: Dict[str, str],
    results: Dict[str, GradeResult],
    overall_grade: str,
    resubmit_questions: List[Tuple[str, str]]
) -> BytesIO:
    """
    Create a Word document with the grading feedback.

    Returns:
        BytesIO object containing the document
    """
    doc = Document()

    # Set up default style
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = BODY_SIZE

    # Title - "Certified SA Pro Plan Building Assignment" on one line
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Certified SA Pro Plan Building Assignment ")
    title_run.bold = True
    title_run.font.name = FONT_NAME
    title_run.font.size = TITLE_SIZE
    title_run.font.color.rgb = PURPLE

    # Header info - all on same line style like Amanda's
    header_para = doc.add_paragraph()

    # Student Name
    label_run = header_para.add_run("Student Name: ")
    label_run.bold = True
    label_run.font.name = FONT_NAME
    label_run.font.color.rgb = LABEL_RED
    value_run = header_para.add_run(f"{student_name} ")
    value_run.font.name = FONT_NAME

    # Submission Date
    label_run = header_para.add_run("Submission Date: ")
    label_run.bold = True
    label_run.font.name = FONT_NAME
    label_run.font.color.rgb = LABEL_RED
    value_run = header_para.add_run(f"{submission_date}")
    value_run.font.name = FONT_NAME

    # Reviewed By and Grade on next line
    header_para2 = doc.add_paragraph()
    label_run = header_para2.add_run("Reviewed By: ")
    label_run.bold = True
    label_run.font.name = FONT_NAME
    label_run.font.color.rgb = LABEL_RED
    value_run = header_para2.add_run(f"{reviewer_name} ")
    value_run.font.name = FONT_NAME

    label_run = header_para2.add_run("Grade: ")
    label_run.bold = True
    label_run.font.name = FONT_NAME
    label_run.font.color.rgb = LABEL_RED
    grade_run = header_para2.add_run(overall_grade)
    grade_run.bold = True
    grade_run.font.name = FONT_NAME
    grade_run.font.color.rgb = RED if overall_grade == "Resubmit" else GREEN

    doc.add_paragraph()  # Space

    # Question definitions
    question_info = {
        'q1': ("Question 1 - Maisie's Plan 1 Target Duration", "Maisie"),
        'q2': ("Question 2 - Maisie's Plan 2 Target Duration", "Maisie"),
        'q3': ("Question 3 - Maisie's Plan 3 Target Duration", "Maisie"),
        'q4': ("Question 4 - Maisie After Struggle", "Maisie"),
        'q5': ("Question 5 - Minna's Plan 1 Target Duration", "Minna"),
        'q6': ("Question 6 - Minna's Plan 2 Target Duration", "Minna"),
        'q7': ("Question 7 - Minna's Plan 3 Target Duration", "Minna"),
        'q8': ("Question 8 - Minna Target Duration Increase", "Minna"),
        'q9': ("Question 9 - Oliver's Plan 1 Target Duration", "Oliver"),
        'q10': ("Question 10 - Oliver's Plan 2 Target Duration", "Oliver"),
        'q11': ("Question 11 - Oliver's Plan 3 Target Duration", "Oliver"),
        'q12': ("Question 12 - Oliver Keys Testing", "Oliver"),
        'q13': ("Question 13 - Bella's Plan 1 Target Duration", "Bella"),
        'q13b': ("Question 13B - Bella's Plan 1 Warmups", "Bella"),
        'q14': ("Question 14 - Bella's Plan 2 Target Duration", "Bella"),
        'q14b': ("Question 14B - Bella's Plan 2 Warmups", "Bella"),
        'q15': ("Question 15 - Bella's Plan 3 Target Duration", "Bella"),
        'q15b': ("Question 15B - Bella's Plan 3 Warmups", "Bella"),
        'q16': ("Question 16 - Bella Car Protocol", "Bella"),
        'q17': ("Question 17 - DIAB Warmups", "DIAB"),
    }

    # Add each question's grading
    for q_id in ['q1', 'q2', 'q3', 'q4', 'q5', 'q6', 'q7', 'q8', 'q9', 'q10', 'q11', 'q12',
                 'q13', 'q13b', 'q14', 'q14b', 'q15', 'q15b', 'q16', 'q17']:

        if q_id not in results:
            continue

        q_title, dog = question_info.get(q_id, (q_id, ""))
        result = results[q_id]
        answer = answers.get(q_id, "")

        # Add question header - Purple like Amanda's
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(q_title)
        q_run.bold = True
        q_run.font.name = FONT_NAME
        q_run.font.color.rgb = PURPLE

        # Add student's answer
        ans_para = doc.add_paragraph()
        ans_label = ans_para.add_run("Your answer: ")
        ans_label.bold = True
        ans_label.font.name = FONT_NAME
        ans_label.font.color.rgb = PURPLE
        ans_value = ans_para.add_run(str(answer) if answer else "(no answer)")
        ans_value.font.name = FONT_NAME
        ans_value.font.color.rgb = PURPLE

        # Add grade - on its own line, colored
        grade_para = doc.add_paragraph()
        grade_text = "CORRECT" if result.is_correct else "INCORRECT"
        grade_run = grade_para.add_run(grade_text)
        grade_run.bold = True
        grade_run.font.name = FONT_NAME
        grade_run.font.color.rgb = GREEN if result.is_correct else RED

        # Add calculation if present (italicized)
        if result.calculation:
            calc_para = doc.add_paragraph()
            calc_run = calc_para.add_run(result.calculation)
            calc_run.italic = True
            calc_run.font.name = FONT_NAME
            calc_run.font.color.rgb = PURPLE

        # Add feedback
        feedback_para = doc.add_paragraph()
        feedback_run = feedback_para.add_run(result.feedback)
        feedback_run.font.name = FONT_NAME
        feedback_run.font.color.rgb = PURPLE

        doc.add_paragraph()  # Space between questions

    # Add overall summary
    doc.add_paragraph()
    summary_heading = doc.add_paragraph()
    summary_run = summary_heading.add_run(f"{student_name},")
    summary_run.bold = True
    summary_run.font.name = FONT_NAME
    summary_run.font.color.rgb = PURPLE

    # Generate summary based on results
    summary_text = generate_summary(results, overall_grade, student_name)
    summary_para = doc.add_paragraph()
    summary_text_run = summary_para.add_run(summary_text)
    summary_text_run.font.name = FONT_NAME
    summary_text_run.font.color.rgb = PURPLE

    # If resubmit, add list of questions to redo
    if resubmit_questions:
        doc.add_paragraph()
        resubmit_para = doc.add_paragraph()
        resubmit_run = resubmit_para.add_run(
            "Review the following questions and send your updated responses directly to me via email:"
        )
        resubmit_run.bold = True

        for q_id, q_label in resubmit_questions:
            q_num = q_id.replace('q', 'Question ').replace('b', 'B')
            doc.add_paragraph(f"- {q_num}: {q_label}", style='List Bullet')

        # Add helpful resources if needed - check for percentage calculation errors
        has_increase_errors = any(q_id in ['q2', 'q3', 'q6', 'q7', 'q10', 'q14', 'q15']
                                  for q_id, _ in resubmit_questions)
        if has_increase_errors:
            doc.add_paragraph()
            doc.add_paragraph(
                "If you haven't seen it, here's an app that's really helpful with the calculations "
                "for duration increases. When using this app put in the time then add the percentage. "
                "Here are examples of how to enter the info into the app. It does all the math for you."
            )
            doc.add_paragraph("1:00 + 10% =")
            doc.add_paragraph(":30 + 20% =")
            doc.add_paragraph()
            link_para = doc.add_paragraph()
            link_para.add_run(
                "https://apps.apple.com/app/id1270645866"
            )

    # Save to BytesIO
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)

    return doc_buffer


def generate_summary(results: Dict[str, GradeResult], overall_grade: str, student_name: str) -> str:
    """Generate a personalized summary paragraph based on grading results."""

    # Check specific areas of strength
    got_diab_right = (results.get('q4', GradeResult(False, '')).is_correct and
                      results.get('q5', GradeResult(False, '')).is_correct)

    got_drops_right = results.get('q11', GradeResult(False, '')).is_correct

    got_keys_right = results.get('q12', GradeResult(False, '')).is_correct

    got_warmups_right = (results.get('q13b', GradeResult(False, '')).is_correct and
                        results.get('q14b', GradeResult(False, '')).is_correct)

    # Check areas needing work
    increase_errors = [q for q in ['q2', 'q3', 'q6', 'q7', 'q10', 'q14', 'q15']
                       if not results.get(q, GradeResult(True, '')).is_correct]

    initial_td_errors = [q for q in ['q1', 'q9', 'q13']
                         if not results.get(q, GradeResult(True, '')).is_correct]

    # Build summary in Amanda's warm, conversational style
    parts = []

    # Opening - vary based on performance
    if overall_grade == "Cleared":
        parts.append("Overall, you've demonstrated a strong understanding of separation anxiety training principles.")
    else:
        parts.append("Overall, you've demonstrated a good understanding of many aspects of separation anxiety training.")

    # Specific praise - written naturally
    strengths = []
    if got_diab_right:
        strengths.append("You correctly identified when to use DIAB for both Maisie and Minna")
    if got_drops_right:
        strengths.append("showed good judgment in dropping Oliver's duration when he struggled")
    if got_keys_right:
        strengths.append("correctly chose to decrease target duration when reintroducing anxiety-provoking cues")
    if got_warmups_right:
        strengths.append("Your warmup management for Bella was excellent, correctly reducing warmups when they caused agitation")

    if strengths:
        # Join naturally
        if len(strengths) == 1:
            parts.append(strengths[0] + ".")
        elif len(strengths) == 2:
            parts.append(strengths[0] + ", and " + strengths[1] + ".")
        else:
            parts.append(strengths[0] + ", " + strengths[1] + ", and " + strengths[2] + "." if len(strengths) == 3 else ", ".join(strengths[:-1]) + ", and " + strengths[-1] + ".")

    # Areas for improvement - friendly tone
    if increase_errors and initial_td_errors:
        parts.append("\n\nThere are just a few questions I want you to review and resubmit. I encourage you to review Module 2 for clarification on percentage increases as well as the body language lessons for setting initial target durations.")
    elif increase_errors:
        parts.append("\n\nWhere you are struggling is how to calculate the math for percentage increases. If you haven't seen it, there's a helpful app for duration calculations - I've included the link below.")
    elif initial_td_errors:
        parts.append("\n\nThere are just a few questions I want you to review and resubmit. I encourage you to review the body language lessons and take another look at setting initial target durations.")

    # Closing - warm and encouraging
    if overall_grade == "Cleared":
        parts.append("\n\nGreat work!")
    else:
        parts.append("\n\nYour strong grasp of the fundamentals and logical reasoning about training decisions shows excellent potential. With attention to these specific guidelines, your plan building skills will be well-rounded.")

    return " ".join(parts)
